"""Procesador de archivos con redacción de datos sensibles.

Soporta PDF (PyMuPDF), Markdown (.md) y Word (.docx).
Reemplaza cada dato sensible con su etiqueta [TIPO].
"""

from __future__ import annotations

import logging
import shutil
import time
from pathlib import Path

import fitz

from app.exceptions import PDFProcessingError
from app.models import DetectedEntity, ProcessingResult
from app.ner.engine import NEREngine

logger = logging.getLogger(__name__)


# Color map for redaction backgrounds (light pastel colors)
_TAG_COLORS: dict[str, tuple[float, float, float]] = {
    "NOMBRE": (1, 0.85, 0.85),
    "EMAIL": (0.85, 0.92, 1),
    "TELEFONO": (0.85, 1, 0.85),
    "CELULAR": (0.85, 1, 0.92),
    "DIRECCION": (1, 0.95, 0.8),
    "DNI": (0.95, 0.85, 1),
    "CUIT_CUIL": (0.9, 0.85, 1),
    "TARJETA_CREDITO": (1, 0.85, 0.92),
    "CUENTA_BANCARIA": (0.85, 0.95, 0.95),
    "PASAPORTE": (1, 1, 0.8),
}


class PDFProcessor:
    """Procesa documentos aplicando redacciones basadas en entidades detectadas.

    Soporta archivos PDF, Markdown (.md) y Word (.docx).

    Attributes:
        output_dir: Directorio donde se guardan los documentos ofuscados.
        ner_engine: Motor NER para detección de entidades sensibles.
    """

    def __init__(self, output_dir: str, ner_engine: NEREngine) -> None:
        """Inicializa el procesador de documentos.

        Args:
            output_dir: Ruta al directorio de salida (e.g. 'ofuscados/').
            ner_engine: Instancia del motor NER configurado.
        """
        self._output_dir = Path(output_dir)
        self._ner_engine = ner_engine
        self._ensure_output_dir()

    def process_file(self, file_path: str) -> ProcessingResult:
        """Procesa un archivo según su extensión.

        Despacha al método apropiado según el tipo de archivo:
        - .pdf → _process_pdf_file
        - .md → _process_markdown_file
        - .docx → _process_docx_file

        Args:
            file_path: Ruta absoluta o relativa al archivo de entrada.

        Returns:
            ProcessingResult con estadísticas del procesamiento.
        """
        input_path = Path(file_path)
        extension = input_path.suffix.lower()

        if extension == ".pdf":
            return self._process_pdf_file(file_path)
        elif extension == ".md":
            return self._process_markdown_file(file_path)
        elif extension == ".docx":
            return self._process_docx_file(file_path)
        else:
            return ProcessingResult(
                input_file=file_path,
                output_file="",
                success=False,
                entities_found=0,
                entities_by_type={},
                error=f"Formato no soportado: {extension}",
            )

    def _process_pdf_file(self, file_path: str) -> ProcessingResult:
        """Procesa un archivo PDF completo.

        Extrae texto por página, detecta entidades sensibles y
        aplica redacciones generando un nuevo PDF en la carpeta
        de salida.

        Args:
            file_path: Ruta absoluta o relativa al PDF de entrada.

        Returns:
            ProcessingResult con estadísticas del procesamiento.
        """
        start_time = time.perf_counter()
        input_path = Path(file_path)
        output_path = self._generate_output_path(file_path)

        try:
            doc = self._open_document(file_path)
        except PDFProcessingError as exc:
            elapsed_ms = int((time.perf_counter() - start_time) * 1000)
            return ProcessingResult(
                input_file=str(input_path),
                output_file=str(output_path),
                success=False,
                entities_found=0,
                entities_by_type={},
                error=exc.message,
                processing_time_ms=elapsed_ms,
            )

        try:
            all_entities: list[DetectedEntity] = []

            # Primera pasada: solo detectar entidades (sin modificar)
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if not text.strip():
                    continue

                entities = self._ner_engine.detect(text, page_num)
                all_entities.extend(entities)

            doc.close()

            if not all_entities:
                # No se encontraron datos sensibles — no generar archivo
                elapsed_ms = int((time.perf_counter() - start_time) * 1000)
                return ProcessingResult(
                    input_file=str(input_path),
                    output_file=str(output_path),
                    success=True,
                    entities_found=0,
                    entities_by_type={},
                    processing_time_ms=elapsed_ms,
                )

            # Copiar el original a la ruta de salida
            self._ensure_output_dir()
            shutil.copy2(str(input_path), str(output_path))

            # Abrir la copia y aplicar redacciones preservando formato
            doc_copy = fitz.open(str(output_path))
            for page_num in range(len(doc_copy)):
                page = doc_copy[page_num]
                page_entities = [e for e in all_entities if e.page == page_num]
                if page_entities:
                    self._apply_redactions(page, page_entities)

            # Guardar in-place (incremental) para preservar estructura del PDF
            doc_copy.saveIncr()
            doc_copy.close()

            # Generar versión Markdown ofuscada
            self._generate_markdown(input_path, all_entities)

            # Calcular estadísticas
            entities_by_type = self._count_entities_by_type(all_entities)
            elapsed_ms = int((time.perf_counter() - start_time) * 1000)

            return ProcessingResult(
                input_file=str(input_path),
                output_file=str(output_path),
                success=True,
                entities_found=len(all_entities),
                entities_by_type=entities_by_type,
                processing_time_ms=elapsed_ms,
            )

        except Exception as exc:
            doc.close()
            elapsed_ms = int((time.perf_counter() - start_time) * 1000)
            error_msg = f"Error procesando PDF: {exc}"
            logger.error(error_msg)
            return ProcessingResult(
                input_file=str(input_path),
                output_file=str(output_path),
                success=False,
                entities_found=0,
                entities_by_type={},
                error=error_msg,
                processing_time_ms=elapsed_ms,
            )

    def _process_markdown_file(self, file_path: str) -> ProcessingResult:
        """Procesa un archivo Markdown.

        Lee el texto, detecta entidades sensibles, reemplaza con
        etiquetas y guarda en ofuscados/ con sufijo _ofuscado.

        Args:
            file_path: Ruta al archivo Markdown de entrada.

        Returns:
            ProcessingResult con estadísticas del procesamiento.
        """
        start_time = time.perf_counter()
        input_path = Path(file_path)
        output_name = f"{input_path.stem}_ofuscado.md"
        output_path = self._output_dir / output_name

        try:
            text = input_path.read_text(encoding="utf-8")

            if not text.strip():
                elapsed_ms = int((time.perf_counter() - start_time) * 1000)
                return ProcessingResult(
                    input_file=str(input_path),
                    output_file=str(output_path),
                    success=True,
                    entities_found=0,
                    entities_by_type={},
                    processing_time_ms=elapsed_ms,
                )

            # Detectar entidades en todo el texto (página 0)
            entities = self._ner_engine.detect(text, page=0)

            if not entities:
                elapsed_ms = int((time.perf_counter() - start_time) * 1000)
                return ProcessingResult(
                    input_file=str(input_path),
                    output_file=str(output_path),
                    success=True,
                    entities_found=0,
                    entities_by_type={},
                    processing_time_ms=elapsed_ms,
                )

            # Aplicar redacciones al texto
            redacted_text = self._apply_text_redactions(text, entities)

            # Guardar en ofuscados/
            self._ensure_output_dir()
            output_path.write_text(redacted_text, encoding="utf-8")

            # Generar versión en ofuscados_md/ también
            self._generate_markdown_from_text(input_path, redacted_text)

            entities_by_type = self._count_entities_by_type(entities)
            elapsed_ms = int((time.perf_counter() - start_time) * 1000)

            return ProcessingResult(
                input_file=str(input_path),
                output_file=str(output_path),
                success=True,
                entities_found=len(entities),
                entities_by_type=entities_by_type,
                processing_time_ms=elapsed_ms,
            )

        except Exception as exc:
            elapsed_ms = int((time.perf_counter() - start_time) * 1000)
            error_msg = f"Error procesando Markdown: {exc}"
            logger.error(error_msg)
            return ProcessingResult(
                input_file=str(input_path),
                output_file=str(output_path),
                success=False,
                entities_found=0,
                entities_by_type={},
                error=error_msg,
                processing_time_ms=elapsed_ms,
            )

    def _process_docx_file(self, file_path: str) -> ProcessingResult:
        """Procesa un archivo Word (.docx).

        Usa python-docx para leer párrafos, detecta entidades sensibles,
        reemplaza texto en los párrafos y guarda en ofuscados/ con
        sufijo _ofuscado.

        Args:
            file_path: Ruta al archivo Word de entrada.

        Returns:
            ProcessingResult con estadísticas del procesamiento.
        """
        start_time = time.perf_counter()
        input_path = Path(file_path)
        output_name = f"{input_path.stem}_ofuscado.docx"
        output_path = self._output_dir / output_name

        try:
            from docx import Document

            doc = Document(str(input_path))

            # Extraer todo el texto de los párrafos
            full_text = "\n".join(para.text for para in doc.paragraphs)

            if not full_text.strip():
                elapsed_ms = int((time.perf_counter() - start_time) * 1000)
                return ProcessingResult(
                    input_file=str(input_path),
                    output_file=str(output_path),
                    success=True,
                    entities_found=0,
                    entities_by_type={},
                    processing_time_ms=elapsed_ms,
                )

            # Detectar entidades en todo el texto
            all_entities = self._ner_engine.detect(full_text, page=0)

            if not all_entities:
                elapsed_ms = int((time.perf_counter() - start_time) * 1000)
                return ProcessingResult(
                    input_file=str(input_path),
                    output_file=str(output_path),
                    success=True,
                    entities_found=0,
                    entities_by_type={},
                    processing_time_ms=elapsed_ms,
                )

            # Aplicar redacciones en cada párrafo del documento
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                # Detectar entidades en este párrafo específico
                para_entities = self._ner_engine.detect(para.text, page=0)
                if para_entities:
                    redacted = self._apply_text_redactions(
                        para.text, para_entities
                    )
                    # Reemplazar el texto del párrafo preservando el primer run
                    if para.runs:
                        para.runs[0].text = redacted
                        for run in para.runs[1:]:
                            run.text = ""
                    else:
                        para.text = redacted

            # Guardar documento ofuscado
            self._ensure_output_dir()
            doc.save(str(output_path))

            # Generar versión Markdown
            redacted_full = self._apply_text_redactions(full_text, all_entities)
            self._generate_markdown_from_text(input_path, redacted_full)

            entities_by_type = self._count_entities_by_type(all_entities)
            elapsed_ms = int((time.perf_counter() - start_time) * 1000)

            return ProcessingResult(
                input_file=str(input_path),
                output_file=str(output_path),
                success=True,
                entities_found=len(all_entities),
                entities_by_type=entities_by_type,
                processing_time_ms=elapsed_ms,
            )

        except Exception as exc:
            elapsed_ms = int((time.perf_counter() - start_time) * 1000)
            error_msg = f"Error procesando Word: {exc}"
            logger.error(error_msg)
            return ProcessingResult(
                input_file=str(input_path),
                output_file=str(output_path),
                success=False,
                entities_found=0,
                entities_by_type={},
                error=error_msg,
                processing_time_ms=elapsed_ms,
            )

    def _open_document(self, file_path: str) -> fitz.Document:
        """Abre un documento PDF manejando errores comunes.

        Args:
            file_path: Ruta al archivo PDF.

        Returns:
            Documento fitz abierto.

        Raises:
            PDFProcessingError: Si el PDF está corrupto o protegido.
        """
        try:
            doc = fitz.open(file_path)
        except Exception as exc:
            msg = f"No se pudo abrir el PDF '{file_path}': {exc}"
            logger.error(msg)
            raise PDFProcessingError(
                code="CORRUPTED",
                message=msg,
                recoverable=True,
            ) from exc

        if doc.is_encrypted:
            doc.close()
            msg = (
                f"El PDF '{file_path}' está protegido con contraseña. "
                "Se omite del procesamiento."
            )
            logger.warning(msg)
            raise PDFProcessingError(
                code="PASSWORD_PROTECTED",
                message=msg,
                recoverable=True,
            )

        return doc

    def _apply_redactions(
        self, page: fitz.Page, entities: list[DetectedEntity]
    ) -> None:
        """Aplica redacciones en una página usando PyMuPDF nativo.

        Para cada entidad detectada, busca el texto en la página,
        marca la zona con add_redact_annot y aplica la redacción
        reemplazando con la etiqueta [TIPO].

        Args:
            page: Página del documento fitz.
            entities: Entidades detectadas en esta página.
        """
        if not entities:
            return

        # Pre-extract all text info from the page for font detection
        page_dict = page.get_text("dict")
        span_map: list[tuple[float, float, float, float, float, str]] = []
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    bbox = span.get("bbox", (0, 0, 0, 0))
                    span_map.append((
                        bbox[0], bbox[1], bbox[2], bbox[3],
                        span.get("size", 11.0),
                        span.get("font", "Helvetica"),
                    ))

        for entity in entities:
            label = f"[{entity.entity_type}]"

            # Estrategia: buscar el texto original de la entidad en la página.
            # Si la entidad fue fusionada (merge), el texto puede no coincidir
            # exactamente con lo que hay en el PDF. En ese caso, buscamos
            # las palabras individuales y creamos un rect que las cubra todas.
            text_instances = page.search_for(entity.text)

            if not text_instances:
                # Si no se encuentra el texto fusionado, buscar por palabras
                # individuales y crear un bounding box que las cubra
                words = entity.text.split()
                if words:
                    # Buscar la primera y última palabra
                    first_rects = page.search_for(words[0])
                    last_rects = page.search_for(words[-1])

                    if first_rects and last_rects:
                        # Tomar el primer match de la primera palabra
                        # y el primer match de la última palabra que esté
                        # en la misma línea (similar y-coordinate)
                        for fr in first_rects:
                            for lr in last_rects:
                                # Verificar que están en la misma línea
                                if abs(fr.y0 - lr.y0) < 5:
                                    # Crear rect que cubra desde primera hasta última
                                    combined = fitz.Rect(
                                        fr.x0, min(fr.y0, lr.y0),
                                        lr.x1, max(fr.y1, lr.y1)
                                    )
                                    text_instances = [combined]
                                    break
                            if text_instances:
                                break

            for rect in text_instances:
                fontsize = 11.0
                fontname = "Helvetica"
                best_overlap = 0.0

                for sx0, sy0, sx1, sy1, ssize, sfont in span_map:
                    ox0 = max(rect.x0, sx0)
                    oy0 = max(rect.y0, sy0)
                    ox1 = min(rect.x1, sx1)
                    oy1 = min(rect.y1, sy1)
                    if ox0 < ox1 and oy0 < oy1:
                        overlap = (ox1 - ox0) * (oy1 - oy0)
                        if overlap > best_overlap:
                            best_overlap = overlap
                            fontsize = ssize
                            fontname = sfont

                redact_font = "helv"
                fl = fontname.lower()
                if "courier" in fl or "mono" in fl or "consol" in fl:
                    redact_font = "cour"
                elif "times" in fl or "serif" in fl:
                    redact_font = "tiro"
                elif "symbol" in fl:
                    redact_font = "symb"

                fill_color = _TAG_COLORS.get(
                    entity.entity_type, (0.9, 0.9, 0.9)
                )

                page.add_redact_annot(
                    rect,
                    text=label,
                    fontsize=fontsize,
                    fontname=redact_font,
                    align=fitz.TEXT_ALIGN_LEFT,
                    fill=fill_color,
                    text_color=(0, 0, 0),
                )

        page.apply_redactions()

    def _generate_output_path(self, input_path: str) -> Path:
        """Genera la ruta de salida con sufijo _ofuscado.

        Ejemplo: mipdf.pdf → ofuscados/mipdf_ofuscado.pdf

        Args:
            input_path: Ruta al archivo de entrada.

        Returns:
            Path al archivo de salida en la carpeta ofuscados/.
        """
        path = Path(input_path)
        stem = path.stem
        ext = path.suffix
        output_name = f"{stem}_ofuscado{ext}"
        return self._output_dir / output_name

    def _save_document(self, doc: fitz.Document, output_path: Path) -> None:
        """Guarda el documento procesado en la ruta de salida.

        Args:
            doc: Documento fitz con redacciones aplicadas.
            output_path: Ruta donde guardar el PDF ofuscado.

        Raises:
            PDFProcessingError: Si no se puede escribir el archivo.
        """
        try:
            self._ensure_output_dir()
            doc.save(str(output_path))
            logger.info("PDF ofuscado guardado en: %s", output_path)
        except Exception as exc:
            msg = f"Error al guardar PDF ofuscado en '{output_path}': {exc}"
            logger.error(msg)
            raise PDFProcessingError(
                code="WRITE_ERROR",
                message=msg,
                recoverable=True,
            ) from exc

    def _ensure_output_dir(self) -> None:
        """Crea el directorio de salida si no existe."""
        self._output_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _count_entities_by_type(
        entities: list[DetectedEntity],
    ) -> dict[str, int]:
        """Cuenta entidades agrupadas por tipo.

        Args:
            entities: Lista de entidades detectadas.

        Returns:
            Diccionario con conteo por tipo (e.g. {"NOMBRE": 3, "DNI": 2}).
        """
        counts: dict[str, int] = {}
        for entity in entities:
            type_name = entity.entity_type
            counts[type_name] = counts.get(type_name, 0) + 1
        return counts

    def _generate_markdown(
        self, input_path: Path, entities: list[DetectedEntity]
    ) -> None:
        """Genera una versión Markdown del PDF con datos ofuscados.

        Extrae el texto del PDF original y reemplaza los datos sensibles
        con sus etiquetas [TIPO], guardando el resultado como .md en
        la carpeta ofuscados_md/.

        Args:
            input_path: Ruta al PDF original.
            entities: Entidades detectadas con posiciones por página.
        """
        try:
            md_output_dir = self._output_dir.parent / "ofuscados_md"
            md_output_dir.mkdir(parents=True, exist_ok=True)

            doc = fitz.open(str(input_path))
            md_lines: list[str] = []
            md_lines.append(f"# {input_path.stem}\n")
            md_lines.append(f"*Archivo original: {input_path.name}*\n")
            md_lines.append("---\n")

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if not text.strip():
                    continue

                page_entities = [e for e in entities if e.page == page_num]
                redacted_text = self._apply_text_redactions(text, page_entities)

                md_lines.append(f"\n## Página {page_num + 1}\n")
                md_lines.append(redacted_text)
                md_lines.append("\n")

            doc.close()

            md_filename = f"{input_path.stem}_ofuscado.md"
            md_path = md_output_dir / md_filename
            md_path.write_text("\n".join(md_lines), encoding="utf-8")
            logger.info("Markdown ofuscado guardado en: %s", md_path)

        except Exception as exc:
            logger.warning(
                "No se pudo generar Markdown para '%s': %s",
                input_path.name, exc,
            )

    def _generate_markdown_from_text(
        self, input_path: Path, redacted_text: str
    ) -> None:
        """Genera una versión Markdown en ofuscados_md/ a partir de texto ya redactado.

        Args:
            input_path: Ruta al archivo original (para nombre).
            redacted_text: Texto ya con redacciones aplicadas.
        """
        try:
            md_output_dir = self._output_dir.parent / "ofuscados_md"
            md_output_dir.mkdir(parents=True, exist_ok=True)

            md_lines: list[str] = [
                f"# {input_path.stem}\n",
                f"*Archivo original: {input_path.name}*\n",
                "---\n",
                "",
                redacted_text,
            ]

            md_filename = f"{input_path.stem}_ofuscado.md"
            md_path = md_output_dir / md_filename
            md_path.write_text("\n".join(md_lines), encoding="utf-8")
            logger.info("Markdown ofuscado guardado en: %s", md_path)

        except Exception as exc:
            logger.warning(
                "No se pudo generar Markdown para '%s': %s",
                input_path.name, exc,
            )

    @staticmethod
    def _apply_text_redactions(
        text: str, entities: list[DetectedEntity]
    ) -> str:
        """Reemplaza datos sensibles en texto plano con etiquetas [TIPO].

        Procesa los reemplazos de atrás hacia adelante para no
        alterar las posiciones de las entidades anteriores.

        Args:
            text: Texto original de la página.
            entities: Entidades detectadas en esta página.

        Returns:
            Texto con datos sensibles reemplazados por etiquetas.
        """
        if not entities:
            return text

        sorted_entities = sorted(entities, key=lambda e: e.start, reverse=True)

        for entity in sorted_entities:
            label = f"[{entity.entity_type}]"
            text = text[:entity.start] + label + text[entity.end:]

        return text
